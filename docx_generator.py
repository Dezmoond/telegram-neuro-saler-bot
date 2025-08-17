import os
from datetime import datetime
from typing import Dict, List
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.shared import OxmlElement, qn

class DocxGenerator:
    def __init__(self, dialogs_docx_folder: str = "dialogs_docx"):
        self.dialogs_docx_folder = dialogs_docx_folder
        if not os.path.exists(dialogs_docx_folder):
            os.makedirs(dialogs_docx_folder)

    def create_dialog_docx(self, user_id: int, dialog_data: Dict) -> str:
        """Создает DOCX файл с историей диалога"""
        doc = Document()
        
        # Настройка стилей
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        style.font.size = Pt(12)
        
        # Заголовок
        title = doc.add_heading('История диалога с нейропродажником', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Информация о диалоге
        info_para = doc.add_paragraph()
        info_para.add_run('ID пользователя: ').bold = True
        info_para.add_run(str(user_id))
        info_para.add_run('\nДата начала: ').bold = True
        info_para.add_run(dialog_data.get('start_time', 'Не указано'))
        info_para.add_run('\nДата окончания: ').bold = True
        info_para.add_run(dialog_data.get('end_time', 'Не указано'))
        info_para.add_run('\nПричина завершения: ').bold = True
        info_para.add_run(dialog_data.get('finish_reason', 'Не указано'))
        info_para.add_run('\nКоличество сообщений: ').bold = True
        info_para.add_run(str(len(dialog_data.get('messages', []))))
        
        doc.add_paragraph()  # Пустая строка
        
        # Диалог
        doc.add_heading('Диалог', level=1)
        
        messages = dialog_data.get('messages', [])
        for i, msg in enumerate(messages, 1):
            # Время сообщения
            timestamp = msg.get('timestamp', '')
            if timestamp:
                time_para = doc.add_paragraph()
                time_para.add_run(f'[{timestamp}]').italic = True
                time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Сообщение клиента
            client_msg = msg.get('client_message', '')
            if client_msg:
                client_para = doc.add_paragraph()
                client_para.add_run('Клиент: ').bold = True
                client_para.add_run(client_msg)
            
            # Ответ нейропродажника
            neuro_response = msg.get('neuro_salesman_response', '')
            if neuro_response:
                neuro_para = doc.add_paragraph()
                neuro_para.add_run('Нейропродажник: ').bold = True
                neuro_para.add_run(neuro_response)
            
            # JSON коммуникация агентов (если есть)
            agent_comm = msg.get('agent_communication', {})
            if agent_comm:
                agent_para = doc.add_paragraph()
                agent_para.add_run('JSON коммуникация агентов: ').bold = True
                agent_para.add_run(str(agent_comm)).italic = True
            
            doc.add_paragraph()  # Пустая строка между сообщениями
        
        # Сохраняем файл
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        filename = f"dialog_{user_id}_{timestamp}.docx"
        filepath = os.path.join(self.dialogs_docx_folder, filename)
        doc.save(filepath)
        
        return filepath

    def get_docx_file_path(self, user_id: int, dialog_data: Dict) -> str:
        """Возвращает путь к DOCX файлу для отправки в Telegram"""
        timestamp = datetime.now().strftime("%Y-%m-%d_%H-%M-%S")
        filename = f"dialog_{user_id}_{timestamp}.docx"
        return os.path.join(self.dialogs_docx_folder, filename)

    def add_feedback_to_docx(self, docx_filepath: str, feedback: str) -> bool:
        """Добавляет отзыв пользователя в существующий DOCX файл"""
        try:
            # Открываем существующий документ
            doc = Document(docx_filepath)
            
            # Добавляем раздел с отзывом в конец документа
            doc.add_paragraph()  # Пустая строка
            doc.add_paragraph()  # Пустая строка
            
            # Заголовок раздела отзыва
            feedback_heading = doc.add_heading('РЕЗУЛЬТАТ ОПРОСА', level=1)
            feedback_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Добавляем отзыв
            feedback_para = doc.add_paragraph()
            feedback_para.add_run('Отзыв пользователя: ').bold = True
            feedback_para.add_run(feedback)
            
            # Добавляем дату отзыва
            feedback_date_para = doc.add_paragraph()
            feedback_date_para.add_run('Дата отзыва: ').bold = True
            feedback_date_para.add_run(datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
            
            # Сохраняем изменения
            doc.save(docx_filepath)
            return True
            
        except Exception as e:
            print(f"Ошибка при добавлении отзыва в DOCX: {e}")
            return False

